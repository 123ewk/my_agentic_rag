"""
多格式文档加载器
支持PDF、Word、Excel、网页、Markdown等格式
"""
# 1. 基础Python标准库
from abc import ABC, abstractmethod       # 定义抽象基类，用来做加载器的统一接口. abstractmethod,被它标记的方法，必须在子类里重写实现，否则代码直接报错。
from typing import List, Dict, Any, Optional  # 类型注解，让代码更清晰
from pathlib import Path                  # 处理文件路径
import json                               # 处理JSON文件

# 2. LangChain 核心
from langchain_core.documents import Document  # LangChain 的文档对象，所有文件最终都要转成这个格式

# 3. 不同文件格式的解析库
from pypdf import PdfReader               # 读取PDF文件（替代旧的PyPDF2）
from docx import Document as DocxDocument # 读取Word (.docx)文件
import pandas as pd                       # 读取Excel/CSV表格文件
from bs4 import BeautifulSoup             # 解析HTML/网页内容
import requests                           # 爬取网页内容

# 导入日志配置
from loguru import logger


class BaseLoader(ABC):
    """
    基础加载器，定义了加载器的统一接口
    """
    @abstractmethod
    def load(self, file_path: str) -> List[Document]:
        """
        加载文件，返回文档列表
        """
        pass

    @abstractmethod
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        提取元数据
        """
        pass

class PDFLoader(BaseLoader):
    """
    PDF 加载器
    """
    def load(self, file_path: str) -> List[Document]:
        """
        加载 PDF 文件，返回文档列表
        """
        pdf_reader = PdfReader(file_path) # 读取PDF文件
        documents = []
        for page_num,page in enumerate(pdf_reader.pages):
            text = page.extract_text()
            if text.strip():
                documents.append(Document(page_content=text, metadata={
                    "page": page_num,
                    "source": file_path,
                    "file_type": "pdf",
                    "total_pages": len(pdf_reader.pages)
                    }))
        return documents
    
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        提取 PDF 文件元数据
        """
        pdf_reader = PdfReader(file_path)
        metadata = pdf_reader.metadata
        return {
            "title": metadata.get("/Title", Path(file_path).stem),  
            "author": metadata.get("/Author", "Unknown"),
            "creation_date": metadata.get("/CreationDate", ""),
            "total_pages": len(pdf_reader.pages)
        }


class WordLoader(BaseLoader):
    """
    Word 加载器
    """
    def load(self, file_path: str) -> List[Document]:
        """
        加载 Word 文件，返回文档列表
        """     
        docx = DocxDocument(file_path)
        texts = []
            
        # 按段落提取
        for para in docx.paragraphs:
            text = para.text.strip()
            if text:
                texts.append(para.text)
        
        # 合并为一个文档（可根据需要调整）
        full_text = "\n".join(texts)
            
        return [Document(
            page_content=full_text,
            metadata={
                "source": file_path,
                "type": "docx",
                "paragraphs": len(texts)
               }
        )]
    
    
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """提取Word元数据"""
        docx = DocxDocument(file_path)
        return {
            "title": docx.core_properties.title or Path(file_path).stem,
            "author": docx.core_properties.author or "Unknown",
            "created": str(docx.core_properties.created),
            "paragraphs": len([p for p in docx.paragraphs if p.text.strip()])
        }

class MarkdownLoader(BaseLoader):
    """
    Markdown 加载器
    """
    def load(self, file_path: str) -> List[Document]:
        """
        加载 Markdown 文件，返回文档列表
        """
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
        return [Document(page_content=text, metadata={
            "source": file_path,
            "type": "markdown",
        })]


    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """提取Markdown元数据(Front Matter)"""
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
            
        # 简单的Front Matter(前置内容)解析
        metadata = {"title": Path(file_path).stem}
        if text.startswith('---'):
            parts = text.split('---', 2) # 两个 --- 之间的 Front Matter 内容
            if len(parts) >= 3: # 确保文件有完整的 Front Matter 结构（开头和结尾都有 ---）
                for line in parts[1].strip().split('\n'):
                    if ':' in line:
                        key, value = line.split(':', 1)
                        metadata[key.strip()] = value.strip()
        
        return metadata
        

class WebLoader(BaseLoader):
    """
    网页加载器
    """
    def load(self, url: str) -> List[Document]:
        """
        加载网页内容
        """
        try:
            response = requests.get(url, timeout=10)
            soup = BeautifulSoup(response.text, 'html.parser')

            # 提取正文（简化版）
            for script in soup(["script", "style"]):
                script.decompose() # 用 .decompose() 把它们从 HTML 中彻底删除，避免这些内容混入正文
            
            text = soup.get_text(separator='\n', strip=True) #separator='\n'：不同元素的文本之间用换行分隔，保留段落结构
            
            return [Document(
                page_content=text,
                metadata={
                    "source": url,
                    "type": "web"
                }
            )]
            
        except Exception as e:
            logger.error(f"加载网页失败: {url}: {e}")
            return []
    
    def extract_metadata(self, url: str) -> Dict[str, Any]:
        """提取网页元数据"""
        return {
            "url": url,
            "type": "web"
        }


class ExcelCSVLoader(BaseLoader):
    """
    Excel 和 CSV 表格加载器
    支持 .xlsx, .xls, .csv 格式
    """
    
    def load(self, file_path: str) -> List[Document]:
        """
        加载 Excel 或 CSV 文件，返回文档列表
        每个 sheet 或 CSV 数据作为一个文档返回
        
        Args:
            file_path: 文件路径
            
        Returns:
            Document 列表，每个 sheet 或 CSV 数据对应一个文档
        """
        path = Path(file_path) # 转换为 Path 对象，方便操作文件路径
        file_suffix = path.suffix.lower()# 获取文件后缀并转换为小写
            
        if file_suffix == '.csv':
            return self._load_csv(file_path)
        elif file_suffix in ['.xlsx', '.xls', '.xlsm']:
            return self._load_excel(file_path)
        else:
            logger.warning(f"不支持的文件格式: {file_path}")
            return []
                
    
    def _load_csv(self, file_path: str) -> List[Document]:
        """
        加载 CSV 文件
        
        Args:
            file_path: CSV 文件路径
            
        Returns:
            包含 CSV 内容的 Document 列表
        """
        df = pd.read_csv(file_path) # CSV 文件本质上就是单个表格，它没有 “工作表（Sheet）” 的概念，一个文件对应一张表。
        return self._dataframe_to_documents(df, file_path, "csv")
    
    def _load_excel(self, file_path: str) -> List[Document]:
        """
        加载 Excel 文件，处理多个 sheet
        
        Args:
            file_path: Excel 文件路径
            
        Returns:
            包含每个 sheet 内容的 Document 列表
        """
        excel_file = pd.ExcelFile(file_path)
        documents = []

        # # Excel 文件（.xlsx/.xls）的核心特点是支持多工作表（Sheet），一个文件里可以存多个表格。
        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(excel_file, sheet_name=sheet_name)
            sheet_docs = self._dataframe_to_documents(
                df, 
                file_path, 
                "excel",
                sheet_name=sheet_name
            )
            documents.extend(sheet_docs)
            
        return documents
    
    def _dataframe_to_documents(
        self, 
        df: pd.DataFrame,
        file_path: str, 
        file_type: str,
        sheet_name: Optional[str] = None
    ) -> List[Document]:
        """
        将 DataFrame 转换为 Document 对象
        
        Args:
            df: pandas DataFrame 对象
            file_path: 源文件路径
            file_type: 文件类型 ('csv' 或 'excel')
            sheet_name: Excel sheet 名称（仅 Excel 文件需要）
            
        Returns:
            Document 列表
        """
        documents = []
        
        # 处理空 DataFrame
        if df.empty:
            return documents
        
        # 将 DataFrame 转换为文本表示
        text_parts = []
        
        # 添加标题行
        if file_type == 'csv' or (file_type == 'excel' and sheet_name):
            text_parts.append(f"=== {sheet_name if sheet_name else '数据'} ===")
        
        # 转换为表格文本格式
        text_parts.append(df.to_string(index=False))
        
        full_text = "\n".join(text_parts)
        
        # 构建元数据
        metadata = {
            "source": file_path,
            "file_type": file_type,
            "rows": len(df),
            "columns": len(df.columns),
            "column_names": list(df.columns)
        }
        
        # 如果是 Excel，添加 sheet 名称
        if sheet_name:
            metadata["sheet_name"] = sheet_name
        
        documents.append(Document(
            page_content=full_text,
            metadata=metadata
        ))
        
        return documents
    
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        提取 Excel/CSV 文件元数据
        
        Args:
            file_path: 文件路径
            
        Returns:
            元数据字典
        """
        try:
            path = Path(file_path)
            file_suffix = path.suffix.lower()
            
            metadata = {
                "title": path.stem,
                "file_type": file_suffix.lstrip('.'),
            }
            
            if file_suffix == '.csv':
                df = pd.read_csv(file_path)
                metadata.update({
                    "rows": len(df),
                    "columns": len(df.columns),
                    "column_names": list(df.columns)
                })
            elif file_suffix in ['.xlsx', '.xls', '.xlsm']:
                excel_file = pd.ExcelFile(file_path)
                metadata["sheet_names"] = excel_file.sheet_names
                metadata["total_sheets"] = len(excel_file.sheet_names)
                
                # 获取第一个 sheet 的基本信息
                if excel_file.sheet_names:
                    df = pd.read_excel(excel_file, sheet_name=excel_file.sheet_names[0])
                    metadata.update({
                        "rows": len(df),
                        "columns": len(df.columns),
                        "column_names": list(df.columns)
                    })
            else:
                logger.warning(f"不支持的文件格式: {file_path}")
                return {}
            
            return metadata
            
        except Exception as e:
            logger.error(f"提取表格文件元数据失败: {file_path}: {e}")
            return {}


class DocumentLoader:
    """
    统一文档加载器，用于加载不同格式的文档文件
    """

    def __init__(self):
        self.loaders = {
            '.pdf': PDFLoader(),
            '.docx': WordLoader(),
            '.md': MarkdownLoader(),
            '.txt': MarkdownLoader(),  # 复用Markdown加载器
            '.csv': ExcelCSVLoader(),
            '.xlsx': ExcelCSVLoader(),
            '.xls': ExcelCSVLoader(),
            '.xlsm': ExcelCSVLoader(),
            '.html': WebLoader(),
            '.htm': WebLoader(),
        }
    
    def load(self, file_path: str) -> List[Document]:
        """根据文件类型选择加载器"""
        path = Path(file_path)
        suffix = path.suffix.lower()
        
        if suffix in self.loaders:
            return self.loaders[suffix].load(file_path)
        else:
            raise ValueError(f"现不支持的文件类型: {suffix}")
    
    def load_batch(self, file_paths: List[str]) -> List[Document]:
        """批量加载文档"""
        all_documents = []
        for file_path in file_paths:
            try:
                docs = self.load(file_path)
                all_documents.extend(docs)
            except Exception as e:
                logger.error(f"加载文件 {file_path} 时出错: {e}")
        
        return all_documents

def get_document_loader() -> DocumentLoader:
    """
    获取文档加载器实例
    
    返回：
        配置好的DocumentLoader实例
    """
    return DocumentLoader()
